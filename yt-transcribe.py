#!/usr/bin/env python3
"""
yt-transcribe.py

CLI utility for downloading YouTube captions and exporting them as plain text
or Word documents. See tmp/goals.md for the full specification guiding this
implementation.
"""

from __future__ import annotations

import argparse
import logging
import os
import re
import shutil
import subprocess
import sys
import time
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Sequence, Tuple
from urllib.parse import parse_qs, urlparse

VIDEO_ID_PATTERN = re.compile(r"^[A-Za-z0-9_-]{11}$")
DEFAULT_LANG_PREF = "en,en-US,en-GB"
STAGE_DIRECTIONS = (
    "music",
    "applause",
    "laughter",
    "silence",
    "inaudible",
)
STAGE_DIRECTION_PATTERN = re.compile(
    r"\[(?:music|applause|laughter|silence|inaudible)\]",
    re.IGNORECASE,
)


class CliError(RuntimeError):
    """Error raised for user-visible problems."""


@dataclass
class TranscriptSelection:
    """Metadata about the chosen transcript."""

    kind: str  # manual, generated, translated
    source_kind: Optional[str]
    language: str
    translated_to: Optional[str]
    entries: List[dict]


def configure_logging(verbose: bool) -> None:
    """Configure the root logger once."""
    logging.basicConfig(level=logging.DEBUG if verbose else logging.INFO, format="%(message)s")
    logging.debug("Verbose logging enabled.")


def expand_language_list(lang_option: str) -> List[str]:
    """Parse the --lang option into a clean list of language codes."""
    if not lang_option:
        return []
    codes = [code.strip() for code in lang_option.split(",")]
    return [code for code in codes if code]


def normalize_url(url_or_id: str) -> str:
    """Ensure URLs missing a scheme are upgraded to https://."""
    if url_or_id.startswith("www."):
        return f"https://{url_or_id}"
    return url_or_id


def extract_video_id(url_or_id: str) -> Tuple[str, str]:
    """Extract the 11-character YouTube video ID and return it with the canonical watch URL."""
    candidate = (url_or_id or "").strip()
    if not candidate:
        raise CliError("You must provide a YouTube URL or 11-character video ID.")

    if VIDEO_ID_PATTERN.fullmatch(candidate):
        video_id = candidate
        logging.debug("Interpreted input as raw video ID: %s", video_id)
        return video_id, f"https://www.youtube.com/watch?v={video_id}"

    normalized = normalize_url(candidate)
    parsed = urlparse(normalized)
    logging.debug("Parsed URL: %s", parsed)
    video_id: Optional[str] = None

    host = parsed.netloc.lower()
    path_segments = [segment for segment in parsed.path.split("/") if segment]

    if host.endswith("youtu.be"):
        if path_segments:
            video_id = path_segments[0]
            logging.debug("Extracted ID from youtu.be URL: %s", video_id)
    elif "youtube.com" in host:
        if parsed.path.startswith("/watch"):
            query = parse_qs(parsed.query)
            video_id = query.get("v", [None])[0]
            logging.debug("Extracted ID from watch URL query: %s", video_id)
        elif path_segments:
            first_segment = path_segments[0]
            if first_segment in {"shorts", "embed", "live"} and len(path_segments) >= 2:
                video_id = path_segments[1]
                logging.debug("Extracted ID from %s path: %s", first_segment, video_id)

    if not video_id:
        raise CliError(f"Unable to locate a valid video ID in input: {url_or_id!r}")

    if not VIDEO_ID_PATTERN.fullmatch(video_id):
        raise CliError(f"Invalid video ID extracted: {video_id!r}")

    canonical_url = f"https://www.youtube.com/watch?v={video_id}"
    return video_id, canonical_url


def probe_ytdlp(skip_probe: bool = False) -> Optional[Path]:
    """Locate a usable yt-dlp binary if available."""
    if skip_probe:
        logging.debug("Skipping yt-dlp probe due to --no-yt-dlp flag.")
        return None

    which_result = shutil.which("yt-dlp")
    if which_result:
        path = Path(which_result)
        logging.debug("Found yt-dlp on PATH: %s", path)
        return path

    for candidate in (Path("/opt/homebrew/bin/yt-dlp"), Path("/usr/local/bin/yt-dlp")):
        if candidate.is_file() and os.access(candidate, os.X_OK):
            logging.debug("Found yt-dlp at fallback location: %s", candidate)
            return candidate

    logging.debug("yt-dlp not found on system.")
    return None


def get_video_title(yt_dlp_path: Path, canonical_url: str, timeout: float = 30.0) -> Optional[str]:
    """Use yt-dlp to retrieve the video title. Returns None on failure."""
    command = [
        str(yt_dlp_path),
        "--no-playlist",
        "--skip-download",
        "--get-title",
        canonical_url,
    ]
    logging.debug("Running yt-dlp for title: %s", " ".join(command))
    try:
        result = subprocess.run(
            command,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=timeout,
        )
    except FileNotFoundError:
        logging.warning("yt-dlp binary disappeared: %s", yt_dlp_path)
        return None
    except subprocess.TimeoutExpired:
        logging.warning("yt-dlp timed out after %.1f seconds while fetching title.", timeout)
        return None
    except subprocess.CalledProcessError as exc:
        stderr = exc.stderr.strip() if exc.stderr else "unknown error"
        logging.warning("yt-dlp failed to fetch title (%s). stderr: %s", exc.returncode, stderr)
        return None

    raw_title = result.stdout.strip()
    if not raw_title:
        logging.warning("yt-dlp returned an empty title for %s", canonical_url)
        return None

    first_line = raw_title.splitlines()[0].strip()
    logging.debug("yt-dlp title resolved: %s", first_line)
    return first_line or None


def sanitize_filename(base: str) -> str:
    """Sanitize a title for filesystem usage."""
    if not base:
        return ""

    text = unicodedata.normalize("NFC", base)

    filtered_chars: List[str] = []
    for char in text:
        category = unicodedata.category(char)
        if category.startswith(("Cc", "Cf")):
            continue  # Drop control characters
        if category.startswith(("Cs", "Co")):
            continue  # Drop private-use / surrogate pairs
        if category == "So":
            continue  # Drop symbols such as emoji/pictographs
        if char in "\\/:*?\"<>|":
            continue
        filtered_chars.append(char)

    sanitized = "".join(filtered_chars)
    sanitized = re.sub(r"\s+", " ", sanitized)
    sanitized = re.sub(r"\s+([,.!?;:])", r"\1", sanitized)
    sanitized = sanitized.strip(" .")

    try:
        from unidecode import unidecode  # type: ignore
    except ImportError:
        sanitized = sanitized.encode("ascii", "ignore").decode("ascii")
    else:
        sanitized = unidecode(sanitized)
    sanitized = sanitized.encode("ascii", "ignore").decode("ascii")
    sanitized = re.sub(r"\s+", " ", sanitized).strip(" .")

    if len(sanitized) > 200:
        sanitized = sanitized[:200].rstrip()

    windows_reserved = {
        "CON",
        "PRN",
        "AUX",
        "NUL",
        *(f"COM{i}" for i in range(1, 10)),
        *(f"LPT{i}" for i in range(1, 10)),
    }
    if sanitized.upper() in windows_reserved:
        sanitized += "_"

    return sanitized


def fetch_transcript(
    video_id: str,
    prefer_generated: bool,
    languages: Sequence[str],
    translate_to: Optional[str] = None,
    retries: int = 2,
) -> TranscriptSelection:
    """Fetch a transcript according to the CLI options."""
    try:
        from youtube_transcript_api import (
            NoTranscriptFound,
            NotTranslatable,
            TranscriptsDisabled,
            TranslationLanguageNotAvailable,
            VideoUnavailable,
            YouTubeTranscriptApi,
            YouTubeTranscriptApiException,
        )
    except ImportError as exc:  # pragma: no cover - user environment issue
        raise CliError(
            "Missing dependency 'youtube-transcript-api'. Install it with pip before running this script."
        ) from exc

    language_preferences = list(dict.fromkeys(languages)) if languages else expand_language_list(DEFAULT_LANG_PREF)
    logging.debug("Language preferences: %s", language_preferences)

    api = YouTubeTranscriptApi()
    transcripts = None
    for attempt in range(retries + 1):
        try:
            transcripts = api.list(video_id)
            break
        except VideoUnavailable as err:
            raise CliError("The video is unavailable or restricted in this region.") from err
        except TranscriptsDisabled as err:
            raise CliError("Captions are disabled for this video.") from err
        except YouTubeTranscriptApiException as err:
            if attempt == retries:
                raise CliError(f"Failed to fetch transcript list: {err}") from err
            delay = min(5.0, 1.5 * (attempt + 1))
            logging.warning(
                "Transcript list retrieval failed (attempt %d/%d): %s; retrying in %.1fs.",
                attempt + 1,
                retries + 1,
                err,
                delay,
            )
            time.sleep(delay)

    if transcripts is None:
        raise CliError("Unable to fetch transcript metadata after multiple attempts.")

    search_order = [
        ("generated", transcripts.find_generated_transcript),
        ("manual", transcripts.find_manually_created_transcript),
    ]
    if not prefer_generated:
        search_order.reverse()

    chosen = None
    chosen_kind = None

    for kind, finder in search_order:
        try:
            chosen = finder(language_preferences)
            chosen_kind = kind
            logging.debug("Selected %s transcript with language %s", kind, chosen.language_code)
            break
        except NoTranscriptFound:
            logging.debug("No %s transcripts for preferred languages: %s", kind, language_preferences)
            continue

    if not chosen:
        raise CliError("No manual or auto captions available for this video.")

    translation_target = translate_to.strip() if translate_to else None
    translated_to = None

    try:
        base_entries = chosen.fetch()
    except Exception as err:  # pragma: no cover - library/network errors
        raise CliError(f"Failed to download transcript entries: {err}") from err

    if translation_target:
        try:
            translated = chosen.translate(translation_target)
            translated_entries = translated.fetch()
            translated_to = translation_target
            logging.info(
                "Using auto-translated captions → %s (source %s)",
                translation_target,
                chosen.language_code,
            )
            language_code = translated.language_code
            normalized_entries = normalize_entries(translated_entries)
        except (NoTranscriptFound, NotTranslatable, TranslationLanguageNotAvailable):
            logging.warning(
                "Translation into %s is unavailable; using original %s captions.",
                translation_target,
                chosen.language_code,
            )
            language_code = chosen.language_code
            normalized_entries = normalize_entries(base_entries)
        except YouTubeTranscriptApiException as err:
            logging.warning(
                "Translation into %s failed (%s); falling back to original captions.",
                translation_target,
                err,
            )
            language_code = chosen.language_code
            normalized_entries = normalize_entries(base_entries)
        else:
            return TranscriptSelection(
                kind="translated",
                source_kind=chosen_kind or "manual",
                language=language_code,
                translated_to=translated_to,
                entries=normalized_entries,
            )
    else:
        language_code = chosen.language_code
        normalized_entries = normalize_entries(base_entries)

    label = "auto-generated" if chosen_kind == "generated" else "manual"
    logging.info("Using %s captions (%s)", label, language_code)

    return TranscriptSelection(
        kind=chosen_kind or "manual",
        source_kind=chosen_kind or "manual",
        language=language_code,
        translated_to=translated_to,
        entries=normalized_entries,
    )


def clean_caption_text(text: str, keep_tags: bool = False) -> str:
    """Apply cleaning rules to caption text."""
    if text is None:
        return ""

    cleaned = unicodedata.normalize("NFC", text)
    if not keep_tags:
        cleaned = STAGE_DIRECTION_PATTERN.sub("", cleaned)

    cleaned = cleaned.replace("\n", " ")
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = re.sub(r"\s+([,.!?;:])", r"\1", cleaned)
    return cleaned.strip()


def normalize_entries(raw_entries: Sequence[object]) -> List[dict]:
    """Convert transcript entries into dictionaries with text/start/duration."""
    normalized: List[dict] = []
    for item in raw_entries:
        if isinstance(item, dict):
            text = item.get("text", "")
            start = item.get("start", 0.0)
            duration = item.get("duration", 0.0)
        else:
            text = getattr(item, "text", "")
            start = getattr(item, "start", 0.0)
            duration = getattr(item, "duration", 0.0)
        normalized.append({"text": text, "start": start, "duration": duration})
    return normalized


def format_timestamp(seconds: float) -> str:
    """Format caption start time as MM:SS or HH:MM:SS."""
    if seconds < 0:
        seconds = 0
    total_seconds = int(seconds)
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def write_txt(
    path: Path,
    transcript: TranscriptSelection,
    canonical_url: str,
    title: str,
    include_timestamps: bool,
    keep_tags: bool,
) -> None:
    """Write the transcript as UTF-8 text."""
    path.parent.mkdir(parents=True, exist_ok=True)

    lines: List[str] = []
    if title:
        lines.append(f"Title: {title}")
    lines.append(f"Source: {canonical_url}")
    lines.append("")

    for entry in transcript.entries:
        raw_text = entry.get("text", "")
        is_stage_direction = False
        if keep_tags and raw_text:
            if STAGE_DIRECTION_PATTERN.fullmatch(raw_text.strip()):
                is_stage_direction = True

        text = clean_caption_text(raw_text, keep_tags=keep_tags)

        if not keep_tags and raw_text and STAGE_DIRECTION_PATTERN.fullmatch(raw_text.strip()):
            if lines and lines[-1] != "":
                lines.append("")
            continue

        if not text:
            continue
        if is_stage_direction and lines and lines[-1] != "":
            lines.append("")
        if include_timestamps:
            timestamp = format_timestamp(entry.get("start", 0.0))
            lines.append(f"{timestamp} - {text}")
        else:
            lines.append(text)

    content = "\n".join(lines).rstrip() + "\n"
    path.write_text(content, encoding="utf-8")


def write_docx(
    path: Path,
    transcript: TranscriptSelection,
    canonical_url: str,
    title: str,
    include_timestamps: bool,
    keep_tags: bool,
) -> None:
    """Write the transcript as a DOCX document."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.opc.constants import RELATIONSHIP_TYPE  # type: ignore
    except ImportError as exc:  # pragma: no cover - user environment issue
        raise CliError(
            "Missing dependency 'python-docx'. Install it with pip before using --docx."
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = document.styles["Normal"]
    font = style.font
    if font.name != "Calibri":
        font.name = "Calibri"
    font.size = Pt(11)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1

    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)
        new_run.append(r_pr)

        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    heading_text = title or "YouTube Transcript"
    heading_para = document.add_heading(level=1)
    if canonical_url:
        add_hyperlink(heading_para, heading_text, canonical_url)
    else:
        heading_para.add_run(heading_text)

    source_paragraph = document.add_paragraph()
    source_paragraph.add_run("Source: ").bold = True
    source_paragraph.add_run(canonical_url)

    document.add_paragraph()  # blank line
    previous_blank = True

    for entry in transcript.entries:
        raw_text = entry.get("text", "")
        is_stage_direction = False
        if keep_tags and raw_text:
            if STAGE_DIRECTION_PATTERN.fullmatch(raw_text.strip()):
                is_stage_direction = True

        text = clean_caption_text(raw_text, keep_tags=keep_tags)

        if not keep_tags and raw_text and STAGE_DIRECTION_PATTERN.fullmatch(raw_text.strip()):
            if not previous_blank:
                document.add_paragraph()
                previous_blank = True
            continue

        if not text:
            continue

        if is_stage_direction and not previous_blank:
            document.add_paragraph()
            previous_blank = True

        paragraph = document.add_paragraph()
        if include_timestamps:
            timestamp = format_timestamp(entry.get("start", 0.0))
            paragraph.add_run(f"{timestamp} - ").bold = True
        paragraph.add_run(text)
        previous_blank = False

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    """Set up the CLI parser and return parsed arguments."""
    parser = argparse.ArgumentParser(
        prog="yt-transcribe.py",
        description="Download YouTube captions (manual or auto) and export to text or DOCX.",
    )
    parser.add_argument("url_or_id", help="YouTube URL (any common form) or 11-char video ID.")
    parser.add_argument("--docx", action="store_true", help="Export as .docx instead of plain text.")
    parser.add_argument("--out", type=str, help="Explicit output path (overrides automatic naming).")
    parser.add_argument("--nostamp", action="store_true", help="Omit timestamps from the output.")
    parser.add_argument("--gencaps", action="store_true", help="Prefer auto-generated captions over manual.")
    parser.add_argument(
        "--lang",
        default="en,en-US,en-GB",
        help="Comma-separated list of preferred caption languages (default: en,en-US,en-GB).",
    )
    parser.add_argument(
        "--translate",
        help=(
            "Request YouTube auto-translation into the given language code "
            "(falls back to original if unavailable)."
        ),
    )
    parser.add_argument("--keep-tags", action="store_true", help="Keep stage directions like [Music].")
    parser.add_argument("--no-yt-dlp", action="store_true", help="Skip yt-dlp title lookup.")
    parser.add_argument("-v", "--verbose", action="store_true", help="Enable verbose logging.")
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)
    configure_logging(args.verbose)
    logging.debug("Arguments: %s", args)

    try:
        languages = expand_language_list(args.lang)
        video_id, canonical_url = extract_video_id(args.url_or_id)
        logging.debug("Canonical URL: %s", canonical_url)

        yt_dlp_path = probe_ytdlp(args.no_yt_dlp)
        if yt_dlp_path:
            logging.debug("Using yt-dlp at %s", yt_dlp_path)
        else:
            logging.debug("yt-dlp not available; falling back to video ID for title.")

        title = video_id
        if yt_dlp_path:
            fetched_title = get_video_title(yt_dlp_path, canonical_url)
            if fetched_title:
                title = fetched_title
            else:
                logging.info("Falling back to video ID for filename/title because yt-dlp title lookup failed.")

        if args.nostamp:
            logging.info("Timestamps disabled (--nostamp).")
        if args.keep_tags:
            logging.info("Stage direction tags will be preserved (--keep-tags).")
        if args.translate:
            logging.info("Requesting auto-translation into %s", args.translate)

        transcript = fetch_transcript(
            video_id=video_id,
            prefer_generated=args.gencaps,
            languages=languages,
            translate_to=args.translate,
        )

        include_timestamps = not args.nostamp

        sanitized_base = sanitize_filename(title)
        if not sanitized_base:
            sanitized_base = video_id

        extension = ".docx" if args.docx else ".txt"

        output_path: Path
        if args.out:
            raw_out = Path(args.out)
            treat_as_directory = False

            if raw_out.exists() and raw_out.is_dir():
                treat_as_directory = True
            elif args.out.endswith(os.sep):
                treat_as_directory = True

            if treat_as_directory:
                target_dir = raw_out
                target_dir.mkdir(parents=True, exist_ok=True)
                output_path = target_dir / f"{sanitized_base}{extension}"
            else:
                output_path = raw_out
                if output_path.suffix.lower() != extension:
                    output_path = output_path.with_suffix(extension)
        else:
            output_path = Path(f"{sanitized_base}{extension}")

        writer = write_docx if args.docx else write_txt
        writer(
            path=output_path,
            transcript=transcript,
            canonical_url=canonical_url,
            title=title,
            include_timestamps=include_timestamps,
            keep_tags=args.keep_tags,
        )

        logging.info("Transcript saved to %s", output_path)
        return 0
    except CliError as err:
        logging.error("%s", err)
        return 1
    except KeyboardInterrupt:
        logging.error("Interrupted.")
        return 130
    except Exception as err:  # pragma: no cover - safeguard
        logging.exception("Unexpected error: %s", err)
        return 1


if __name__ == "__main__":
    sys.exit(main())
